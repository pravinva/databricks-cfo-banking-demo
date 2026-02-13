#!/usr/bin/env python3
"""
Generate a Databricks POV document for Treasury Modeling (Deposits + Stress + PPNR).

This script creates a .docx deliverable intended for executive/ALCO audiences, with
clear guidance on which model to use when (Approach 1/2/3 + PPNR) and a narrative
about time-to-value enabled by Unity Catalog governance + a unified lakehouse.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


@dataclass(frozen=True)
class ModelRow:
    model: str
    purpose: str
    when_to_use: str
    inputs: str
    outputs: str
    cadence: str


def _add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{key}: ")
    r1.bold = True
    p.add_run(value)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for i in items:
        doc.add_paragraph(i, style="List Bullet")


def _set_compact_table(table) -> None:
    # Reduce default spacing a bit for dense POV tables.
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)


def build_document() -> Document:
    doc = Document()

    # Title page
    title = doc.add_paragraph("Treasury Modeling on Databricks")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Deposits (Approach 1/2/3) + Stress Testing + PPNR")
    subtitle.style = doc.styles["Subtitle"]
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().add_run("")  # spacer
    _add_kv(doc, "Date", date.today().isoformat())
    _add_kv(doc, "Audience", "Treasurer / ALCO, Treasury Analytics, Model Risk, Data & AI Platform")
    _add_kv(doc, "Operating principle", "Batch-first (weekly/monthly refresh). No 24/7 model serving required.")

    doc.add_page_break()

    # Executive POV
    doc.add_heading("1) Executive POV", level=1)
    doc.add_paragraph(
        "Treasury teams need fast, defensible answers to three questions: "
        "(1) how sensitive are deposits to rate changes, "
        "(2) how quickly do balances run off under different conditions, and "
        "(3) what does that mean for earnings (PPNR) under baseline and stress scenarios."
    )

    doc.add_paragraph("This POV focuses on two outcomes:")
    _add_bullets(
        doc,
        [
            "A clear model catalog: which model to use when (Approach 1 vs 2 vs 3, plus PPNR).",
            "Faster time-to-value: unified data + governance + training and refresh on real bank data.",
        ],
    )

    # Fragmented tooling narrative
    doc.add_heading("2) The problem without Databricks: fragmented tooling and slow refresh", level=1)
    doc.add_paragraph(
        "In many banks, treasury modeling is distributed across vendor systems, spreadsheets, "
        "ad‑hoc Python/R environments, and separate BI layers. The outcome is a brittle workflow "
        "that is hard to reproduce, slow to update, and difficult to govern."
    )
    doc.add_paragraph("Typical symptoms:")
    _add_bullets(
        doc,
        [
            "Multiple versions of “the same” deposit balance / segment definitions across teams.",
            "Manual handoffs (extracts → spreadsheets → slides) and limited auditability.",
            "Slow model refresh cycles (quarterly or ad-hoc) because feature engineering and training are not industrialized.",
            "Scenario runs are expensive and inconsistent because data, assumptions, and code are not centralized.",
        ],
    )

    # Databricks POV: what changes
    doc.add_heading("3) What changes with Databricks: one lakehouse, governed data + models + BI", level=1)
    doc.add_paragraph(
        "Databricks reduces time-to-value by putting data engineering, ML, and analytics on a single platform. "
        "Unity Catalog (UC) provides governed access, discoverability, and lineage across tables, volumes, "
        "and registered models—so teams can move from prototype to production without re-platforming."
    )
    doc.add_paragraph("Key accelerators:")
    _add_bullets(
        doc,
        [
            "Unity Catalog for consistent definitions, permissions, and lineage across the full pipeline.",
            "Delta tables as durable, queryable model inputs/outputs (deposit snapshots, runoff, stress results, PPNR).",
            "MLflow model registry with champion/challenger aliases to promote or roll back safely.",
            "Batch scoring and scheduled jobs to refresh the portfolio view weekly/monthly.",
            "UC Volumes for governed report outputs (HTML/PDF) without relying on DBFS FUSE.",
        ],
    )

    # Model catalog
    doc.add_heading("4) Model catalog: which model to use when", level=1)
    doc.add_paragraph(
        "The demo organizes deposit behavior into three independent approaches. "
        "Banks can adopt them in any order based on the business question and governance requirements."
    )

    rows: list[ModelRow] = [
        ModelRow(
            model="Approach 1 — Static Deposit Beta",
            purpose="Operational sensitivity baseline and segmentation (actionable repricing/defend list).",
            when_to_use=(
                "Use for a stable, explainable baseline. Best for weekly/monthly refresh and "
                "day-to-day ALCO decisions where interpretability matters."
            ),
            inputs=(
                "deposit_accounts_historical, deposit_accounts (label), yield_curves\n"
                "Key features: balances, 30d avg balance, rate gap/spread, txn activity, segment/product encodings."
            ),
            outputs="deposit_beta_predictions (account-level predicted_beta)",
            cadence="Weekly / monthly batch scoring",
        ),
        ModelRow(
            model="Approach 2 — Vintage / Survival + Component Decay",
            purpose="Runoff curves by cohort/tenure for liquidity planning (LCR/NSFR-style behavior).",
            when_to_use=(
                "Use when treasury needs behavioral maturity views: how retention changes by vintage, tenure, and segment. "
                "Required for runoff forecasting rather than just “beta today”."
            ),
            inputs="deposit_accounts_historical (vintages), cohort definitions, segment classifications",
            outputs="cohort_survival_rates, deposit_runoff_forecasts, deposit_beta_training_phase2 (enhanced features)",
            cadence="Monthly refresh (or aligned to planning cadence)",
        ),
        ModelRow(
            model="Approach 3 — Dynamic Beta + Stress Testing",
            purpose="Regulatory-style scenario analysis (CCAR-style) with regime-conditional behavior.",
            when_to_use=(
                "Use for scenario narratives and tail risk: rapid hikes/cuts, regime shifts, "
                "and multi-horizon projections. Suitable for stress testing packages."
            ),
            inputs="deposit_beta_training_phase2 (from Approach 2), rate scenarios, portfolio composition",
            outputs="stress_test_results (scenario outputs), dynamic beta parameters (if persisted)",
            cadence="Quarterly / on-demand scenario runs; ad-hoc for shocks",
        ),
        ModelRow(
            model="PPNR (Non‑Interest Income + Non‑Interest Expense → PPNR)",
            purpose="Translate deposit strategy and scenarios into earnings impact (the ALCO “so what”).",
            when_to_use=(
                "Use when decisions must be expressed in earnings terms: "
                "how runoff and pricing affect fee income, expense leverage, and overall PPNR."
            ),
            inputs="general_ledger (training), macro/rate context, treasury scenario inputs",
            outputs="ppnr_forecasts (monthly series)",
            cadence="Monthly refresh; align to forecast cycle",
        ),
    ]

    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "Purpose"
    hdr[2].text = "When to use"
    hdr[3].text = "Inputs (governed tables)"
    hdr[4].text = "Outputs (data products)"
    hdr[5].text = "Refresh cadence"

    for r in rows:
        c = table.add_row().cells
        c[0].text = r.model
        c[1].text = r.purpose
        c[2].text = r.when_to_use
        c[3].text = r.inputs
        c[4].text = r.outputs
        c[5].text = r.cadence

    _set_compact_table(table)

    doc.add_paragraph(
        "Guidance: start with Approach 1 to operationalize segmentation; add Approach 2 for runoff curves; "
        "use Approach 3 when stress narratives are required; integrate PPNR once the earnings storyline is needed."
    )

    # Operating model and governance
    doc.add_heading("5) Operating model (batch-first) and governance", level=1)
    doc.add_paragraph("Recommended cadence (example):")
    _add_bullets(
        doc,
        [
            "Weekly (Sunday night): batch score deposit beta portfolio using the current champion model alias.",
            "Monthly: refresh vintage/runoff curves and PPNR series aligned to planning close.",
            "On-demand: run stress scenarios for rapid hikes/cuts (Approach 3) and publish results to governed tables.",
        ],
    )
    doc.add_paragraph("Governance controls enabled by Unity Catalog + MLflow:")
    _add_bullets(
        doc,
        [
            "Access control and separation of duties: treasury vs model risk vs engineering.",
            "Lineage from source tables → features → models → predictions → reports.",
            "Champion/challenger promotion via aliases (zero-code deployment) and one-click rollback.",
            "Repeatable scoring with deterministic feature definitions (tables act as the contract).",
        ],
    )

    # Time-to-value
    doc.add_heading("6) Time-to-value: why UC matters for treasury", level=1)
    doc.add_paragraph(
        "Time-to-value improves when teams can train and update models on real data without re-wiring "
        "permissions, pipelines, and environments. Unity Catalog provides a consistent namespace and "
        "governance layer so each model can be updated and consumed by downstream reporting and apps "
        "with minimal friction."
    )
    _add_bullets(
        doc,
        [
            "Fewer handoffs: data, code, and model artifacts live in one governed platform.",
            "Faster iteration: update features or training windows and re-train without creating new silos.",
            "Reusable outputs: tables like deposit_beta_predictions, stress_test_results, ppnr_forecasts become data products.",
        ],
    )

    # Appendix
    doc.add_page_break()
    doc.add_heading("Appendix A — Canonical objects (Unity Catalog)", level=1)
    _add_bullets(
        doc,
        [
            "Inputs: cfo_banking_demo.bronze_core_banking.deposit_accounts, deposit_accounts_historical",
            "Rates: cfo_banking_demo.silver_treasury.yield_curves",
            "Approach 1 outputs: cfo_banking_demo.ml_models.deposit_beta_training_data, deposit_beta_predictions",
            "Approach 2 outputs: cfo_banking_demo.ml_models.cohort_survival_rates, deposit_runoff_forecasts, deposit_beta_training_phase2",
            "Approach 3 outputs: cfo_banking_demo.ml_models.stress_test_results",
            "PPNR outputs: cfo_banking_demo.ml_models.ppnr_forecasts",
            "Reports volume: /Volumes/cfo_banking_demo/gold_finance/reports/",
        ],
    )

    return doc


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_path = repo_root / "Databricks_Treasury_Modeling_POV.docx"
    doc = build_document()
    doc.save(out_path.as_posix())
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()

